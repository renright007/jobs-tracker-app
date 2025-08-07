import streamlit as st
import pandas as pd
from datetime import datetime
import os
from PyPDF2 import PdfReader
from docx import Document
from utils import get_db_connection, init_db, ensure_directories
from database_utils import delete_document, save_documents_to_database, migrate_existing_data
from streamlit_shadcn_ui import tabs

# Note: Database initialization is handled by main app.py to prevent cloud filesystem issues
# ensure_directories() also moved to main app initialization

# Run migration to ensure preferred_resume column exists
migrate_existing_data()

# Initialize session state
if 'documents' not in st.session_state:
    st.session_state.documents = pd.DataFrame(columns=[
        'document_name', 'document_type', 'upload_date', 'file_path'
    ])

if 'user_profile' not in st.session_state:
    st.session_state.user_profile = {
        'selected_resume': None,
        'career_goals': "",
        'created_date': None,
        'last_updated_date': None
    }

if 'career_goals' not in st.session_state:
    st.session_state.career_goals = pd.DataFrame(columns=[
        'goals', 'submission_date'
    ])

# Add custom CSS for the container
st.markdown("""
    <style>
    .document-container {
        border: 1px solid #e2e8f0;
        border-radius: 0.5rem;
        padding: 1.5rem;
        background-color: white;
        box-shadow: 0 1px 3px 0 rgb(0 0 0 / 0.1), 0 1px 2px -1px rgb(0 0 0 / 0.1);
    }
    </style>
""", unsafe_allow_html=True)

def extract_text_from_uploaded_file(uploaded_file):
    """Extract text content from uploaded file for database storage."""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        # Reset file pointer to beginning
        uploaded_file.seek(0)
        
        if file_extension == 'pdf':
            # Extract text from PDF
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
            
        elif file_extension == 'docx':
            # Extract text from DOCX
            doc = Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
            
        elif file_extension == 'txt':
            # Read text file
            text = uploaded_file.read().decode('utf-8')
            return text.strip()
            
        else:
            return f"Unsupported file type: {file_extension}"
            
    except Exception as e:
        return f"Error extracting content: {str(e)}"

def upload_document_with_content(uploaded_file, document_name, document_type, user_id):
    """Upload document and store content in database for cloud compatibility."""
    from database_utils import use_supabase
    
    try:
        # Extract text content from uploaded file
        document_content = extract_text_from_uploaded_file(uploaded_file)
        
        if document_content.startswith("Error") or document_content.startswith("Unsupported"):
            return False, document_content
        
        # Prepare document data
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        if use_supabase():
            from supabase_utils import get_supabase_client
            supabase = get_supabase_client()
            
            document_data = {
                'user_id': user_id,
                'document_name': document_name,
                'document_type': document_type,
                'document_content': document_content,
                'upload_date': current_time,
                'file_path': None,  # Cloud-friendly: no file path dependency
                'preferred_resume': 0
            }
            
            result = supabase.table('documents').insert(document_data).execute()
            return True, "Document uploaded and content stored successfully!"
            
        else:
            # SQLite fallback
            conn = get_db_connection()
            c = conn.cursor()
            
            c.execute('''INSERT INTO documents 
                        (user_id, document_name, document_type, document_content, 
                         upload_date, file_path, preferred_resume)
                        VALUES (?, ?, ?, ?, ?, ?, ?)''',
                     (user_id, document_name, document_type, document_content,
                      current_time, None, 0))
            
            conn.commit()
            conn.close()
            return True, "Document uploaded and content stored successfully!"
            
    except Exception as e:
        return False, f"Error uploading document: {str(e)}"


def show_user_portal():
    """Show the User Portal with shadcn tabs."""
    # Authentication is now handled at the main app level
    user_id = st.session_state.get('user_id')
    from database_utils import use_supabase
    
    # Initialize Supabase client once for entire function if using cloud database
    supabase = None
    if use_supabase():
        from supabase_utils import get_supabase_client
        supabase = get_supabase_client()
    
    st.title("User Portal")
    st.info("👋 Manage your documents, set preferences, and track your career goals.")
    
    # Create shadcn tabs with default tab
    selected_tab = tabs(["User Profile", "Document Portal"], default_value="User Profile")
    
    if selected_tab == "User Profile":
        # Load documents from database (filtered by user)
        if use_supabase():
            # Get documents
            docs_result = supabase.table('documents').select('*').eq('document_type', 'Resume').eq('user_id', user_id).execute()
            documents_df = pd.DataFrame(docs_result.data) if docs_result.data else pd.DataFrame()
            
            # Get user profile
            profile_result = supabase.table('user_profile').select('*').eq('user_id', user_id).order('id', desc=True).limit(1).execute()
            profile_df = pd.DataFrame(profile_result.data) if profile_result.data else pd.DataFrame()
        else:
            conn = get_db_connection()
            documents_df = pd.read_sql_query("SELECT * FROM documents WHERE document_type = 'Resume' AND user_id = ?", conn, params=(user_id,))
            profile_df = pd.read_sql_query("SELECT * FROM user_profile WHERE user_id = ? ORDER BY id DESC LIMIT 1", conn, params=(user_id,))
            conn.close()
        
        # Initialize variables with default values
        selected_resume = None
        
        # Load existing profile data if it exists
        if not profile_df.empty:
            selected_resume = profile_df['selected_resume'].iloc[0]
        
        # Simplified Preferred Resume Display
        st.markdown("### 🎯 Current Preferred Resume")
        
        # Get current preferred resume from database
        if use_supabase():
            preferred_result = supabase.table('documents').select('document_name, upload_date').eq('user_id', user_id).eq('preferred_resume', 1).eq('document_type', 'Resume').execute()
            preferred_df = pd.DataFrame(preferred_result.data) if preferred_result.data else pd.DataFrame()
        else:
            conn = get_db_connection()
            preferred_df = pd.read_sql_query(
                "SELECT document_name, upload_date FROM documents WHERE user_id = ? AND preferred_resume = 1 AND document_type = 'Resume'", 
                conn, params=(user_id,)
            )
            conn.close()
        
        if not preferred_df.empty:
            preferred_name = preferred_df['document_name'].iloc[0]
            preferred_date = preferred_df['upload_date'].iloc[0]
            st.success(f"✅ **{preferred_name}** (uploaded: {preferred_date})")
            st.info("💡 This resume will be used by default for all AI-powered features like cover letter generation and job matching.")
            
            # Get the full document data including content
            if use_supabase():
                full_doc_result = supabase.table('documents').select('*').eq('user_id', user_id).eq('preferred_resume', 1).eq('document_type', 'Resume').execute()
                full_doc_df = pd.DataFrame(full_doc_result.data) if full_doc_result.data else pd.DataFrame()
            else:
                conn = get_db_connection()
                full_doc_df = pd.read_sql_query(
                    "SELECT * FROM documents WHERE user_id = ? AND preferred_resume = 1 AND document_type = 'Resume'", 
                    conn, params=(user_id,)
                )
                conn.close()
            
            if not full_doc_df.empty:
                resume_data = full_doc_df.iloc[0]
                
                # Display resume content from database (cloud-friendly)
                with st.expander("👁️ View Resume Content"):
                    if pd.notna(resume_data['document_content']) and resume_data['document_content']:
                        # Use database content (preferred method)
                        st.text_area(
                            "Resume Content (from database)", 
                            value=resume_data['document_content'], 
                            height=300,
                            disabled=True
                        )
                        
                        # Show content stats
                        content_length = len(resume_data['document_content'])
                        word_count = len(resume_data['document_content'].split())
                        st.caption(f"📊 Content Stats: {content_length:,} characters, ~{word_count:,} words")
                        
                    elif pd.notna(resume_data['file_path']) and resume_data['file_path']:
                        # Fallback to file reading for legacy documents
                        st.warning("⚠️ Using legacy file-based content (may not work in cloud)")
                        selected_resume_path = resume_data['file_path']
                        file_extension = selected_resume_path.split('.')[-1].lower()
                        
                        try:
                            # Check if file exists first
                            if not os.path.exists(selected_resume_path):
                                st.error(f"❌ File not found: {selected_resume_path}")
                                st.info("🔧 This is a cloud deployment issue where local files aren't available.")
                                st.info("💡 **Tip:** Re-upload your document to store content in database.")
                                text = None
                            else:
                                if file_extension == 'pdf':
                                    reader = PdfReader(selected_resume_path)
                                    text = ""
                                    for page in reader.pages:
                                        text += page.extract_text() + "\n"
                                elif file_extension == 'docx':
                                    doc = Document(selected_resume_path)
                                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                                elif file_extension == 'txt':
                                    with open(selected_resume_path, 'r', encoding='utf-8') as file:
                                        text = file.read()
                                else:
                                    st.error(f"Unsupported file type: {file_extension}")
                                    text = None
                            
                            if text:
                                st.text_area("Resume Content (from file)", text, height=300, disabled=True)
                                
                        except Exception as e:
                            st.error(f"Error processing file: {str(e)}")
                            st.info("🔧 This might be due to file corruption or cloud deployment limitations.")
                    else:
                        st.warning("⚠️ No content available for this resume.")
                        st.info("💡 Try re-uploading your resume to store content in the database.")
            else:
                st.warning("⚠️ No preferred resume found.")
        else:
            st.warning("⚠️ No preferred resume selected")
            st.info("📋 Go to the **Document Portal** tab to select your preferred resume by checking the box next to it.")
        
        st.markdown("---")
        st.caption("💡 **Tip:** To change your preferred resume, go to Document Portal → Manage Your Documents → Check the 'Preferred Resume' box")
        
        # Career Goals Section
        st.markdown("---")
        st.subheader("What are you looking for?")
        
        # Load career goals from database (filtered by user)
        if use_supabase():
            goals_result = supabase.table('career_goals').select('*').eq('user_id', user_id).order('submission_date', desc=True).limit(1).execute()
            goals_df = pd.DataFrame(goals_result.data) if goals_result.data else pd.DataFrame()
        else:
            conn = get_db_connection()
            goals_df = pd.read_sql_query("SELECT * FROM career_goals WHERE user_id = ? ORDER BY submission_date DESC LIMIT 1", conn, params=(user_id,))
            conn.close()
        
        # Initialize career goals
        career_goals = ""
        if not goals_df.empty:
            career_goals = goals_df['goals'].iloc[0]
        
        # Career Goals input
        career_goals = st.text_area(
            "Describe your career goals, preferred roles, industries, and what you're looking for in your next position...",
            value=career_goals,
            height=300
        )
        
        # Save Career Goals button
        if st.button("Save Career Goals"):
            if career_goals.strip():
                current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                
                if use_supabase():
                    goals_data = {
                        'goals': career_goals,
                        'submission_date': current_time,
                        'user_id': user_id
                    }
                    supabase.table('career_goals').insert(goals_data).execute()
                else:
                    conn = get_db_connection()
                    c = conn.cursor()
                    c.execute('''INSERT INTO career_goals 
                                (goals, submission_date, user_id)
                                VALUES (?, ?, ?)''',
                             (career_goals, current_time, user_id))
                    conn.commit()
                    conn.close()
                
                st.success("Career goals saved successfully!")
                st.rerun()
            else:
                st.warning("Please enter your career goals before saving.")
        
        # Display previous career goals submissions (filtered by user)
        if use_supabase():
            all_goals_result = supabase.table('career_goals').select('*').eq('user_id', user_id).order('submission_date', desc=True).limit(2).execute()
            all_goals_df = pd.DataFrame(all_goals_result.data) if all_goals_result.data else pd.DataFrame()
        else:
            conn = get_db_connection()
            all_goals_df = pd.read_sql_query("SELECT * FROM career_goals WHERE user_id = ? ORDER BY submission_date DESC LIMIT 2", conn, params=(user_id,))
            conn.close()
        
        if not all_goals_df.empty and len(all_goals_df) > 1:
            st.subheader("Previous Submission")
            with st.expander(f"Submission from {all_goals_df.iloc[1]['submission_date']}"):
                st.write(all_goals_df.iloc[1]['goals'])
    
    elif selected_tab == "Document Portal":
        
        # Upload new document section
        st.subheader("📄 Upload New Document")
        uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'docx', 'txt'])
        if uploaded_file is not None:
            document_name = st.text_input("Document Name")
            document_type = st.selectbox("Document Type", ["Resume", "Cover Letter", "Other"])
            
            if st.button("Save Document"):
                if not document_name.strip():
                    st.error("Please enter a document name.")
                else:
                    success, message = upload_document_with_content(uploaded_file, document_name, document_type, user_id)
                    if success:
                        st.success(message)
                        st.rerun()
                    else:
                        st.error(message)
        
        st.markdown("---")
        
        # Document management section
        st.subheader("📋 Manage Your Documents")
        st.info("💡 Check the 'Preferred Resume' box for the resume you want AI features to use by default. Only one can be selected.")
        
        # Load documents from database (include document_content for viewing)
        if use_supabase():
            docs_result = supabase.table('documents').select('id, document_name, document_type, upload_date, preferred_resume, document_content').eq('user_id', user_id).order('upload_date', desc=True).execute()
            documents_df = pd.DataFrame(docs_result.data) if docs_result.data else pd.DataFrame()
        else:
            conn = get_db_connection()
            documents_df = pd.read_sql_query(
                "SELECT id, document_name, document_type, upload_date, preferred_resume, document_content FROM documents WHERE user_id = ? ORDER BY upload_date DESC", 
                conn, params=(user_id,)
            )
            conn.close()
        
        if not documents_df.empty:
            # Convert preferred_resume to boolean for display
            documents_df['preferred_resume'] = documents_df['preferred_resume'].astype(bool)
            
            # Create editable data editor like jobs table
            with st.form("documents_form"):
                st.write("**Edit your documents:**")
                
                # Configure column settings
                column_config = {
                    "id": st.column_config.NumberColumn("ID", disabled=True, width="small"),
                    "document_name": st.column_config.TextColumn("Document Name", disabled=True, width="medium"),
                    "document_type": st.column_config.TextColumn("Type", disabled=True, width="small"),
                    "upload_date": st.column_config.TextColumn("Upload Date", disabled=True, width="medium"),
                    "preferred_resume": st.column_config.CheckboxColumn(
                        "Preferred Resume",
                        help="Check to use this resume for AI features",
                        width="small"
                    )
                }
                
                edited_documents_df = st.data_editor(
                    documents_df,
                    column_config=column_config,
                    hide_index=True,
                    use_container_width=True,
                    height=400,
                    key="documents_editor"
                )
                
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.form_submit_button("💾 Save Changes", use_container_width=True):
                        success, message = save_documents_to_database(user_id, edited_documents_df)
                        if success:
                            st.success(message)
                            st.rerun()
                        else:
                            st.error(message)
                
                with col2:
                    if st.form_submit_button("🗑️ Delete Selected", use_container_width=True):
                        # Get selected rows for deletion (this would need additional logic)
                        st.info("Select documents and use individual delete buttons for now")
            
            # Download section
            st.markdown("---")
            st.subheader("📥 Download Documents")
            for _, row in documents_df.iterrows():
                col1, col2, col3, col4 = st.columns([3, 2, 2, 2])
                
                with col1:
                    status_icon = "⭐" if row['preferred_resume'] else ""
                    st.write(f"{status_icon} **{row['document_name']}**")
                
                with col2:
                    st.write(row['document_type'])
                
                with col3:
                    # Get file path for download
                    if use_supabase():
                        file_path_result = supabase.table('documents').select('file_path').eq('id', row['id']).execute()
                        file_path_df = pd.DataFrame(file_path_result.data) if file_path_result.data else pd.DataFrame()
                    else:
                        conn = get_db_connection()
                        file_path_df = pd.read_sql_query(
                            "SELECT file_path FROM documents WHERE id = ?", 
                            conn, params=(row['id'],)
                        )
                        conn.close()
                    
                    if not file_path_df.empty and file_path_df['file_path'].iloc[0]:
                        file_path = file_path_df['file_path'].iloc[0]
                        if os.path.exists(file_path):
                            with open(file_path, 'rb') as file:
                                st.download_button(
                                    label="📥 Download",
                                    data=file.read(),
                                    file_name=file_path.split('/')[-1],
                                    key=f"download_{row['id']}"
                                )
                        else:
                            st.write("File not found")
                    else:
                        st.write("No file")
                
                with col4:
                    if st.button("🗑️ Delete", key=f"delete_{row['id']}", help="Delete this document"):
                        success, message = delete_document(user_id, row['id'])
                        if success:
                            st.success(message)
                            st.rerun()
                        else:
                            st.error(message)
            # Document Content Viewer Section
            st.markdown("---")
            st.subheader("👁️ View Document Content")
            
            # Document selector for viewing content
            if not documents_df.empty:
                doc_names = [f"{row['document_name']} ({row['document_type']})" for _, row in documents_df.iterrows()]
                doc_ids = documents_df['id'].tolist()
                
                selected_doc_index = st.selectbox(
                    "Select document to view:",
                    range(len(doc_names)),
                    format_func=lambda x: doc_names[x]
                )
                
                if selected_doc_index is not None:
                    selected_doc_id = doc_ids[selected_doc_index]
                    selected_doc = documents_df[documents_df['id'] == selected_doc_id].iloc[0]
                    
                    # Display document content from database
                    with st.expander(f"📄 {selected_doc['document_name']} Content", expanded=True):
                        if pd.notna(selected_doc['document_content']) and selected_doc['document_content']:
                            st.text_area(
                                "Document Content (stored in database)",
                                value=selected_doc['document_content'],
                                height=400,
                                disabled=True,
                                key=f"content_viewer_{selected_doc_id}"
                            )
                            
                            # Show content stats
                            content_length = len(selected_doc['document_content'])
                            word_count = len(selected_doc['document_content'].split())
                            st.caption(f"📊 Content Stats: {content_length:,} characters, ~{word_count:,} words")
                            
                        else:
                            st.warning("⚠️ No content stored in database for this document.")
                            st.info("💡 This document was uploaded before content storage was implemented. Try re-uploading to store content.")
                            
        else:
            st.info("📝 No documents uploaded yet. Upload your first document above!")